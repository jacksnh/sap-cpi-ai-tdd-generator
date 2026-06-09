from openai import OpenAI
import requests
import streamlit as st
from docx import Document
from docx.shared import Inches
from io import BytesIO
import zipfile
import pandas as pd
from datetime import date
from PIL import Image, ImageDraw, ImageFont
import os

# -----------------------------
# OpenRouter Client
# -----------------------------

client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=st.secrets["OPENROUTER_API_KEY"]
)


# -----------------------------
# Streamlit Page Setup
# -----------------------------

st.set_page_config(page_title="SAP CPI TDD Generator", layout="wide")

st.title("SAP CPI Technical Design Document Generator")

st.write("Generate a TDD manually or by uploading an SAP CPI iFlow package and mapping sheet.")


# -----------------------------
# User Inputs
# -----------------------------

generation_mode = st.radio(
    "Choose TDD generation option",
    [
        "Manual Entry - Fetch iFlow using CPI API",
        "Generate using iFlow Package",
        "Generate using iFlow Package + Mapping Sheet"
    ]
)

interface_name = st.text_input("Interface Name")
source_system = st.text_input("Source System")
target_system = st.text_input("Target System")
business_context = st.text_area("Business Context")
business_process = st.text_area("Business Process")

integration_type = st.selectbox(
    "Integration Type",
    ["A2A", "B2B", "To be confirmed"]
)

data_classification = st.selectbox(
    "Data Security Classification",
    ["Public", "Internal", "Confidential", "Strictly Confidential", "To be confirmed"]
)

business_critical = st.selectbox(
    "Business Critical",
    ["Yes", "No", "To be confirmed"]
)

data_transfer_time = st.selectbox(
    "Data Transfer Time",
    ["Synchronous", "<= 2 minutes", "<= 30 minutes", "> 30 minutes", "To be confirmed"]
)

sender_adapter = ""
receiver_adapter = ""
mapping_logic = ""
error_handling = ""
description = ""
cpi_package_name = ""
cpi_iflow_name = ""

if generation_mode == "Manual Entry - Fetch iFlow using CPI API":
    sender_adapter = st.text_input("Sender Adapter", placeholder="HTTPS, SFTP, SOAP, IDoc, etc.")
    receiver_adapter = st.text_input("Receiver Adapter", placeholder="OData, REST, SOAP, SFTP, etc.")
    mapping_logic = st.text_area("Mapping Logic")
    error_handling = st.text_area("Error Handling")
    description = st.text_area("Integration Flow Description")
    st.subheader("Fetch iFlow from CPI API")
    cpi_package_name = st.text_input("CPI Package Name")
    cpi_iflow_name = st.text_input("CPI iFlow Name")

iflow_file = None
mapping_file = None
iflow_summary = ""
mapping_summary = ""
palette_summary = ""

if generation_mode in ["Generate using iFlow Package", "Generate using iFlow Package + Mapping Sheet"]:
    iflow_file = st.file_uploader("Upload SAP CPI iFlow Package ZIP", type=["zip"])

if generation_mode == "Generate using iFlow Package + Mapping Sheet":
    mapping_file = st.file_uploader("Upload Mapping Sheet", type=["xlsx", "xls", "csv"])


# -----------------------------
# Extract iFlow Package Details
# -----------------------------

def extract_iflow_details(zip_bytes):
    details = []

    try:
        zip_bytes.seek(0)

        with zipfile.ZipFile(zip_bytes, "r") as zip_ref:
            file_names = zip_ref.namelist()

            details.append("Files found in iFlow package:")

            for name in file_names:
                details.append(f"- {name}")

            for name in file_names:
                lower_name = name.lower()

                if lower_name.endswith(".iflw") or lower_name.endswith(".xml"):
                    try:
                        content = zip_ref.read(name).decode("utf-8", errors="ignore")
                        details.append(f"\nContent extracted from {name}:")
                        details.append(content[:8000])
                    except Exception:
                        details.append(f"Unable to read {name}")

                if lower_name.endswith(".groovy"):
                    try:
                        content = zip_ref.read(name).decode("utf-8", errors="ignore")
                        details.append(f"\nGroovy script found: {name}")
                        details.append(content[:4000])
                    except Exception:
                        details.append(f"Unable to read {name}")

                if lower_name.endswith(".xsl") or lower_name.endswith(".xslt"):
                    try:
                        content = zip_ref.read(name).decode("utf-8", errors="ignore")
                        details.append(f"\nXSLT mapping found: {name}")
                        details.append(content[:4000])
                    except Exception:
                        details.append(f"Unable to read {name}")

                if lower_name.endswith(".wsdl") or lower_name.endswith(".xsd"):
                    try:
                        content = zip_ref.read(name).decode("utf-8", errors="ignore")
                        details.append(f"\nSchema/WSDL found: {name}")
                        details.append(content[:4000])
                    except Exception:
                        details.append(f"Unable to read {name}")

                if lower_name.endswith(".properties") or lower_name.endswith(".prop"):
                    try:
                        content = zip_ref.read(name).decode("utf-8", errors="ignore")
                        details.append(f"\nProperties file found: {name}")
                        details.append(content[:4000])
                    except Exception:
                        details.append(f"Unable to read {name}")

        return "\n".join(details)

    except Exception as e:
        return f"Error reading iFlow package: {e}"


def extract_iflow_palette_steps_from_zip(zip_bytes):
    steps = []

    try:
        zip_bytes.seek(0)

        with zipfile.ZipFile(zip_bytes, "r") as zip_ref:
            file_names = zip_ref.namelist()

            for name in file_names:
                lower_name = name.lower()

                if lower_name.endswith(".iflw") or lower_name.endswith(".xml"):
                    content = zip_ref.read(name).decode("utf-8", errors="ignore")

                    keywords = [
                        "Content Modifier",
                        "Message Mapping",
                        "Groovy Script",
                        "XSLT Mapping",
                        "Request Reply",
                        "Router",
                        "Filter",
                        "Splitter",
                        "Gather",
                        "Join",
                        "Multicast",
                        "Process Call",
                        "Local Integration Process",
                        "Exception Subprocess",
                        "Start Event",
                        "End Event",
                        "Timer Start",
                        "HTTPS",
                        "HTTP",
                        "SOAP",
                        "OData",
                        "SFTP",
                        "Mail",
                        "JMS",
                        "IDoc",
                        "XI"
                    ]

                    for keyword in keywords:
                        keyword_position = content.lower().find(keyword.lower())

                        if keyword_position != -1:
                            start = max(0, keyword_position - 1000)
                            end = min(len(content), keyword_position + 2000)
                            snippet = content[start:end]

                            steps.append(
                                f"""
Step / Palette Identified: {keyword}
Source File: {name}
Relevant iFlow XML Context:
{snippet}
"""
                            )

        if not steps:
            return "No CPI palette steps could be identified from the iFlow package."

        return "\n\n".join(steps)

    except Exception as e:
        return f"Error extracting iFlow palette steps: {e}"
# -----------------------------
# Extract Mapping Sheet
# -----------------------------

def extract_mapping_sheet(uploaded_file):
    try:
        if uploaded_file.name.endswith(".csv"):
            df = pd.read_csv(uploaded_file)
        else:
            df = pd.read_excel(uploaded_file)

        preview = df.head(50).to_string(index=False)

        return df, preview

    except Exception as e:
        return None, f"Error reading mapping sheet: {e}"


# -----------------------------
# Create Architecture Diagram
# -----------------------------

def create_architecture_diagram(
    interface_name,
    source_system,
    target_system,
    sender_adapter,
    receiver_adapter
):
    width = 1500
    height = 850

    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    try:
        title_font = ImageFont.truetype("arial.ttf", 42)
        box_font = ImageFont.truetype("arial.ttf", 22)
        small_font = ImageFont.truetype("arial.ttf", 16)
        summary_font = ImageFont.truetype("arial.ttf", 36)
        note_font = ImageFont.truetype("arial.ttf", 15)
    except Exception:
        title_font = ImageFont.load_default()
        box_font = ImageFont.load_default()
        small_font = ImageFont.load_default()
        summary_font = ImageFont.load_default()
        note_font = ImageFont.load_default()

    title = f"{source_system or 'Source System'} Integration with {target_system or 'Target System'}"
    draw.text((420, 50), title, fill="black", font=title_font)

    source_box = (170, 150, 370, 390)
    middle_box = (650, 150, 850, 390)
    target_box = (1130, 150, 1330, 390)

    source_color = "#B7E3F0"
    middle_color = "#F6C399"
    target_color = "#B7E3F0"

    draw.rectangle(source_box, fill=source_color, outline="gray", width=2)
    draw.rectangle(middle_box, fill=middle_color, outline="gray", width=2)
    draw.rectangle(target_box, fill=target_color, outline="gray", width=2)

    source_label = source_system or "Source System"
    target_label = target_system or "Target System"
    middleware_label = "SAP CPI /\nIntegration Suite"

    draw.multiline_text((220, 245), source_label, fill="black", font=box_font, spacing=5, align="center")
    draw.multiline_text((685, 235), middleware_label, fill="black", font=box_font, spacing=5, align="center")
    draw.multiline_text((1185, 245), target_label, fill="black", font=box_font, spacing=5, align="center")

    draw.text((235, 205), "SAP", fill="white", font=small_font)
    draw.rectangle((220, 190, 300, 230), fill="#2E9AFE")
    draw.text((235, 202), "SAP", fill="white", font=small_font)

    draw.ellipse((720, 190, 780, 250), fill="#2E86DE")
    draw.text((708, 255), "CPI", fill="black", font=small_font)

    draw.rectangle((1210, 195, 1260, 235), fill="white")
    draw.ellipse((1220, 200, 1250, 210), fill="gray")
    draw.rectangle((1220, 205, 1250, 225), fill="gray")
    draw.ellipse((1220, 220, 1250, 230), fill="gray")

    def arrow_with_label(start, end, label):
     x1, y1 = start
     x2, y2 = end

     draw.line((x1, y1, x2, y2), fill="black", width=3)

     if x2 > x1:
        draw.polygon(
            [(x2, y2), (x2 - 12, y2 - 7), (x2 - 12, y2 + 7)],
            fill="black"
        )
     else:
        draw.polygon(
            [(x2, y2), (x2 + 12, y2 - 7), (x2 + 12, y2 + 7)],
            fill="black"
        )

     mid_x = int((x1 + x2) / 2)
     mid_y = int((y1 + y2) / 2)

     text_box_width = 160
     text_box_height = 22

     draw.rectangle(
        (
            mid_x - text_box_width // 2,
            mid_y - 28,
            mid_x + text_box_width // 2,
            mid_y - 6
        ),
        fill="white"
    )

     draw.text(
        (mid_x - 65, mid_y - 27),
        label,
        fill="black",
        font=small_font
     )

    # Arrows between Source and CPI
    arrow_with_label((370, 230), (650, 230), "Request / Send Data")
    arrow_with_label((650, 275), (370, 275), "Response")

# Arrows between CPI and Target
    arrow_with_label((850, 230), (1130, 230), "Request / Send Data")
    arrow_with_label((1130, 275), (850, 275), "Response")

# Adapter labels
    draw.text((455, 310), sender_adapter or "Sender Adapter", fill="black", font=small_font)
    draw.text((925, 310), receiver_adapter or "Receiver Adapter", fill="black", font=small_font)

    draw.rectangle((180, 470, 420, 480), fill="gray")
    draw.text((170, 495), "Summary", fill="black", font=summary_font)

    summary_text = (
        f"1. {interface_name or 'This interface'} enables integration between "
        f"{source_system or 'source system'} and {target_system or 'target system'} "
        f"using SAP CPI / SAP Integration Suite."
    )

    draw.text((170, 560), summary_text, fill="black", font=small_font)

    note_box = (170, 650, 1330, 735)
    draw.rectangle(note_box, fill="#C9C9C9", outline="#C9C9C9")

    note_text = (
        "Note:\n"
        "1. Arrow direction represents which system initiates the call.\n"
        "2. Black Arrow - Async push call.\n"
        "3. Green Arrow - Sync call.\n"
        "4. Red Arrow - Async pull call."
    )

    draw.multiline_text((180, 660), note_text, fill="black", font=note_font, spacing=4)

    diagram_file = "architecture_diagram.png"
    image.save(diagram_file)

    return diagram_file


# -----------------------------
# Call OpenRouter AI
# -----------------------------

def call_ai(prompt):
    try:
        response = client.chat.completions.create(
            model="openai/gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an SAP CPI Integration Architect. "
                        "Generate accurate technical design documents. "
                        "Do not invent missing information."
                    )
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.2
        )

        return response.choices[0].message.content

    except Exception as e:
        return f"Error calling OpenRouter: {e}"
    

def get_cpi_access_token():
    response = requests.post(
        st.secrets["CPI_TOKEN_URL"],
        data={"grant_type": "client_credentials"},
        auth=(st.secrets["CPI_CLIENT_ID"], st.secrets["CPI_CLIENT_SECRET"])
    )

    if response.status_code != 200:
        raise Exception(f"Token request failed: {response.status_code} - {response.text}")

    return response.json()["access_token"]


def get_package_id_by_name(package_name, token):
    url = f"{st.secrets['CPI_BASE_URL']}/api/v1/IntegrationPackages"

    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/json"
    }

    response = requests.get(url, headers=headers)

    if response.status_code != 200:
        raise Exception(f"Package lookup failed: {response.status_code} - {response.text}")

    packages = response.json().get("d", {}).get("results", [])

    for package in packages:
        if package.get("Name", "").strip().lower() == package_name.strip().lower():
            return package.get("Id")

    raise Exception(f"Package not found: {package_name}")


def get_iflow_id_by_name(package_id, iflow_name, token):
    url = (
        f"{st.secrets['CPI_BASE_URL']}/api/v1/IntegrationPackages"
        f"('{package_id}')/IntegrationDesigntimeArtifacts"
    )

    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/json"
    }

    response = requests.get(url, headers=headers)

    if response.status_code != 200:
        raise Exception(f"iFlow lookup failed: {response.status_code} - {response.text}")

    artifacts = response.json().get("d", {}).get("results", [])

    for artifact in artifacts:
        if artifact.get("Name", "").strip().lower() == iflow_name.strip().lower():
            return artifact.get("Id")

    raise Exception(f"iFlow not found: {iflow_name}")


def download_iflow_from_cpi_by_names(package_name, iflow_name):
    token = get_cpi_access_token()

    package_id = get_package_id_by_name(package_name, token)
    artifact_id = get_iflow_id_by_name(package_id, iflow_name, token)

    url = (
        f"{st.secrets['CPI_BASE_URL']}/api/v1/IntegrationDesigntimeArtifacts"
        f"(Id='{artifact_id}',Version='active')/$value"
    )

    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/zip"
    }

    response = requests.get(url, headers=headers)

    if response.status_code != 200:
        raise Exception(f"iFlow download failed: {response.status_code} - {response.text}")

    return BytesIO(response.content)

# Converts AI-generated Markdown text into proper Word formatting.
# Handles headings (#, ##, ###), Markdown tables, bullet points, numbered lists,
# and removes basic Markdown symbols like **bold** and `code`.

def add_markdown_to_doc(doc, content):
    lines = content.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Headings
        if line.startswith("##### "):
            doc.add_heading(line.replace("##### ", "").strip(), level=5)
            i += 1
            continue

        if line.startswith("#### "):
            doc.add_heading(line.replace("#### ", "").strip(), level=4)
            i += 1
            continue

        if line.startswith("### "):
            doc.add_heading(line.replace("### ", "").strip(), level=3)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line.replace("## ", "").strip(), level=2)
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line.replace("# ", "").strip(), level=1)
            i += 1
            continue

        # Markdown table
        if "|" in line and i + 1 < len(lines) and "---" in lines[i + 1]:
            table_lines = []

            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1

            # Remove separator line like |---|---|
            table_lines = [
                row for row in table_lines
                if not all(cell.strip().replace("-", "") == "" for cell in row.strip("|").split("|"))
            ]

            if table_lines:
                rows = []
                for row in table_lines:
                    cells = [cell.strip() for cell in row.strip("|").split("|")]
                    rows.append(cells)

                max_cols = max(len(row) for row in rows)

                table = doc.add_table(rows=1, cols=max_cols)
                table.style = "Table Grid"

                # Header row
                for col_index in range(max_cols):
                    table.rows[0].cells[col_index].text = rows[0][col_index] if col_index < len(rows[0]) else ""

                # Data rows
                for row in rows[1:]:
                    row_cells = table.add_row().cells
                    for col_index in range(max_cols):
                        row_cells[col_index].text = row[col_index] if col_index < len(row) else ""

            continue

        # Bullet points
        if line.startswith("- "):
            clean_bullet = (
                line.replace("- ", "", 1)
                    .replace("**", "")
                    .replace("__", "")
                    .replace("`", "")
                    .replace("*", "")
            )

            doc.add_paragraph(clean_bullet, style="List Bullet")
            i += 1
            continue

        # Numbered list
        if len(line) > 2 and line[0].isdigit() and line[1] == ".":
            doc.add_paragraph(line, style="List Number")
            i += 1
            continue

        # Normal paragraph
        clean_line = (
    line.replace("**", "")
        .replace("__", "")
        .replace("`", "")
        .replace("*", "")
    )

        doc.add_paragraph(clean_line)
        i += 1
# -----------------------------
# Create Word Document
# -----------------------------

def create_word_document(content, mapping_df=None):
    doc = Document()

    doc.add_heading("SAP CPI Technical Design Document", 0)
    doc.add_paragraph("FORUM Template Style")
    doc.add_heading("1. Interface Details", level=1)
    doc.add_paragraph(f"Interface Name: {interface_name}")
    doc.add_paragraph(f"Source System: {source_system}")
    doc.add_paragraph(f"Target System: {target_system}")
    doc.add_paragraph(f"Sender Adapter: {sender_adapter or 'To be confirmed'}")
    doc.add_paragraph(f"Receiver Adapter: {receiver_adapter or 'To be confirmed'}")

    doc.add_heading("2. Architecture Diagram", level=1)

    try:
        diagram_file = create_architecture_diagram(
            interface_name,
            source_system,
            target_system,
            sender_adapter,
            receiver_adapter
        )
        doc.add_picture(diagram_file, width=Inches(6.5))
    except Exception as e:
        doc.add_paragraph(f"Architecture diagram could not be generated: {e}")

    doc.add_page_break()

    doc.add_heading("Generated SAP CPI Technical Design Document", level=1)
    add_markdown_to_doc(doc, content)

    if mapping_df is not None:
        doc.add_page_break()
        doc.add_heading("Appendix A - Mapping Sheet", level=1)

        table = doc.add_table(rows=1, cols=len(mapping_df.columns))
        table.style = "Table Grid"

        header_cells = table.rows[0].cells

        for i, column in enumerate(mapping_df.columns):
            header_cells[i].text = str(column)

        for _, row in mapping_df.head(100).iterrows():
            row_cells = table.add_row().cells

            for i, value in enumerate(row):
                row_cells[i].text = "" if pd.isna(value) else str(value)

        if len(mapping_df) > 100:
            doc.add_paragraph(
                "Note: Only first 100 rows of mapping sheet are included in this document."
            )

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream


# -----------------------------
# Generate TDD Button
# -----------------------------

if st.button("Generate TDD"):
    mapping_df = None

    if not interface_name or not source_system or not target_system:
        st.warning("Please enter Interface Name, Source System, and Target System.")

    else:
               # Read iFlow either from CPI API or uploaded ZIP
        try:
            if generation_mode == "Manual Entry - Fetch iFlow using CPI API":
                if cpi_package_name and cpi_iflow_name:
                    with st.spinner("Connecting to SAP CPI and downloading iFlow..."):
                        cpi_iflow_zip = download_iflow_from_cpi_by_names(
                            cpi_package_name,
                            cpi_iflow_name
                        )

                        cpi_iflow_zip.seek(0)
                        iflow_summary = extract_iflow_details(cpi_iflow_zip)
                        cpi_iflow_zip.seek(0)
                        palette_summary = extract_iflow_palette_steps_from_zip(cpi_iflow_zip)

                    st.success("iFlow downloaded successfully from SAP CPI.")

                else:
                    st.warning("Please enter CPI Package Name and CPI iFlow Name.")
                    st.stop()

            elif iflow_file is not None:
                with st.spinner("Reading uploaded iFlow package..."):
                    iflow_file.seek(0)
                    zip_bytes = BytesIO(iflow_file.read())
                    iflow_summary = extract_iflow_details(iflow_file)
                    zip_bytes.seek(0)
                    palette_summary = extract_iflow_palette_steps_from_zip(zip_bytes)
        except Exception as e:
            st.error(f"Unable to process iFlow package: {e}")
            st.stop()

        # Mapping sheet upload remains same
        if mapping_file is not None:
            with st.spinner("Reading mapping sheet..."):
                mapping_df, mapping_summary = extract_mapping_sheet(mapping_file)
        current_date = date.today().strftime("%d.%m.%Y")

        prompt = f"""
You are an SAP CPI Integration Architect.

Generate a detailed Technical Design Document for the following SAP CPI interface.

Interface Name: {interface_name}
Source System: {source_system}
Target System: {target_system}
Document Date: {current_date}
Business Context: {business_context}
Business Process: {business_process}
Integration Type: {integration_type}
Data Security Classification: {data_classification}
Business Critical: {business_critical}
Data Transfer Time: {data_transfer_time}

Manual Details:
Sender Adapter: {sender_adapter}
Receiver Adapter: {receiver_adapter}
Mapping Logic: {mapping_logic}
Error Handling: {error_handling}
Integration Flow Description: {description}

Extracted iFlow Package Details:
{iflow_summary}

Mapping Sheet Details:
{mapping_summary}

Extracted CPI Palette / Step Details:
{palette_summary}

Create the document in the same style as a SAP CPI Technical Design Document / FORUM template.

Use the following structure:

1. Introduction
1.1 Business Context
1.2 Business Process
1.3 Detailed Business Requirements

2. Development and Configuration Guidelines

3. Document Change Control / History
3.1 Change History
3.2 Sign Off
3.3 Related Documents
3.4 Issue Log

4. Technical Approach
4.1 Realization Approach
4.1.1 Scope
4.1.2 Out of Scope
4.1.3 Constraint
4.1.4 Risk

4.2 Technical Architecture
4.3 Technical Data Model and Structure
4.3.1 Data Structure
4.3.2 Technical Data Model
4.3.3 Mapping
4.3.4 Routing
4.3.5 Hardware Requirements
4.3.5.1 Used System

5. Realization
5.1 Realization Objects
5.1.1 Realisation Detailed Object List including Sequence and Dependencies
5.1.2 Traceability
5.2 Screen Design
5.3 Security Enablement
5.4 Test Script

6. Cutover Plan for Implementation

7. Appendix
7.1 Glossary
7.2 Meeting Minutes

For each section:
- Write in a formal SAP technical design document style.
- Use tables where appropriate.
- If information is missing, write "To be confirmed".
- Do not invent system names, URLs, credentials, certificates, ports, or personal data.
- Use iFlow package and mapping sheet details where available.
- Mention SAP CPI / SAP Integration Suite as the middleware.

For each identified CPI palette step, explain:
- Step name
- Where it appears in the iFlow sequence
- What input it receives from the previous step
- What operation it performs in this specific interface
- What output it passes to the next step
- Which source/target fields, headers, properties, mappings, scripts, routes, or adapters it is related to
- Why this step is needed in this specific iFlow
- Do not write generic CPI purpose only
- If exact behavior is not available from the iFlow XML, write "To be confirmed from iFlow configuration"

Important rules:
- Use the iFlow package details where available.
- Use the mapping sheet details where available.
- Only use information provided by the user, iFlow package, or mapping sheet.
- Do not invent missing information.
- If any information is missing, mention "To be confirmed".
- Do not include passwords, secrets, API keys, or credentials.
- For Change History, use Document Date as the Date. Do not create random dates.
- Use Version 1.0 and Author as "Integration Team" unless provided by user.
- In Technical Architecture, explain each CPI palette step based on what it is doing in this specific uploaded iFlow. Do not provide generic SAP CPI component descriptions.
"""

        with st.spinner("Generating Technical Design Document..."):
            tdd_content = call_ai(prompt)

        st.subheader("Generated TDD")
        st.write(tdd_content)

        # Generate and show architecture diagram in Streamlit
        try:
            diagram_file = create_architecture_diagram(
                interface_name,
                source_system,
                target_system,
                sender_adapter,
                receiver_adapter
            )

            st.subheader("Architecture Diagram")
            st.image(diagram_file)

            with open(diagram_file, "rb") as img:
                st.download_button(
                    label="Download Architecture Diagram",
                    data=img,
                    file_name="architecture_diagram.png",
                    mime="image/png"
                )

        except Exception as e:
            st.error(f"Architecture diagram could not be generated: {e}")

        # Optional review/edit before download
        edited_tdd = st.text_area(
            "Review/Edit Generated TDD before download",
            value=tdd_content,
            height=500
        )

        word_file = create_word_document(edited_tdd, mapping_df)

        st.download_button(
            label="Download TDD as Word Document",
            data=word_file,
            file_name=f"{interface_name}_TDD.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
